from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.54)

# ── Style Setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    hs.font.name = 'Calibri'
    if i == 1:
        hs.font.size = Pt(20)
    elif i == 2:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(13)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_bold_paragraph(doc, text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p

def add_info_box(doc, text, label="INFO"):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    r = p.add_run(f"[{label}] ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    r2 = p.add_run(text)
    r2.font.italic = True

def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BINA NUSANTARA UNIVERSITY")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Fakultas Ilmu Komputer")
r.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MATERI PERSIAPAN UTS")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SOFTWARE SECURITY")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

doc.add_paragraph()

topics = [
    "1. Threat Modeling (Study Case)",
    "2. Brainstorming (Scenario Analysis, Premortem, Movie Plotting)",
    "3. Attack Tree",
    "4. Chess Strategy"
]
for t in topics:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t)
    r.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Semester 4 — 2025/2026")
r.font.size = Pt(12)
r.font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DAFTAR ISI
# ══════════════════════════════════════════════════════════════════
doc.add_heading('DAFTAR ISI', level=1)
toc_items = [
    ("BAB 1: Threat Modeling (Study Case)", "3"),
    ("   1.1 Definisi Threat Modeling", "3"),
    ("   1.2 Tujuan Threat Modeling", "3"),
    ("   1.3 Metodologi: STRIDE, DREAD, PASTA", "3"),
    ("   1.4 Langkah-Langkah Threat Modeling", "4"),
    ("   1.5 Studi Kasus: Sistem E-Commerce", "4"),
    ("   1.6 Data Flow Diagram (DFD)", "5"),
    ("BAB 2: Brainstorming", "6"),
    ("   2.1 Definisi Brainstorming", "6"),
    ("   2.2 Scenario Analysis", "6"),
    ("   2.3 Premortem Analysis", "7"),
    ("   2.4 Movie Plotting", "8"),
    ("BAB 3: Attack Tree", "10"),
    ("   3.1 Definisi Attack Tree", "10"),
    ("   3.2 Struktur & Notasi", "10"),
    ("   3.3 Contoh Attack Tree", "11"),
    ("   3.4 Kegunaan & Cara Membuat", "11"),
    ("BAB 4: Chess Strategy", "12"),
    ("   4.1 Konsep Chess Strategy", "12"),
    ("   4.2 Analogi Catur-Security", "12"),
    ("   4.3 Prinsip-Prinsip", "13"),
    ("   4.4 Penerapan", "14"),
    ("BAB 5: Soal dan Jawaban", "15"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15))
    r = p.add_run(item)
    r.font.size = Pt(11)
    if not item.startswith("   "):
        r.bold = True
    r2 = p.add_run(f"\t{page}")
    r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# BAB 1: THREAT MODELING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('BAB 1: THREAT MODELING (Study Case)', level=1)

doc.add_heading('1.1 Definisi Threat Modeling', level=2)
doc.add_paragraph(
    'Threat Modeling adalah proses sistematis untuk mengidentifikasi, mengevaluasi, dan memprioritaskan potensi ancaman keamanan terhadap suatu sistem perangkat lunak. '
    'Tujuannya adalah menemukan kerentanan (vulnerability) sebelum penyerang (attacker) mengeksploitasinya.'
)

doc.add_heading('1.2 Tujuan Threat Modeling', level=2)
tujuan = [
    'Mengidentifikasi aset yang perlu dilindungi',
    'Menemukan ancaman yang mungkin menyerang sistem',
    'Menentukan kerentanan yang dapat dieksploitasi',
    'Memprioritaskan risiko berdasarkan dampak dan kemungkinan',
    'Merancang mitigasi yang tepat dan efektif'
]
for t in tujuan:
    doc.add_paragraph(t, style='List Number')

doc.add_heading('1.3 Metodologi Threat Modeling', level=2)

# STRIDE
doc.add_heading('a) STRIDE (Microsoft)', level=3)
doc.add_paragraph(
    'STRIDE adalah kerangka kerja klasifikasi ancaman yang dikembangkan oleh Microsoft. Setiap huruf merepresentasikan satu kategori ancaman:'
)
add_table(doc,
    ['Kategori', 'Deskripsi', 'Contoh'],
    [
        ['Spoofing', 'Menyamar sebagai entitas lain', 'Pemalsuan identitas login'],
        ['Tampering', 'Mengubah data tanpa izin', 'Modifikasi data transaksi'],
        ['Repudiation', 'Menyangkal tindakan yang dilakukan', 'User menyangkal telah melakukan transaksi'],
        ['Information Disclosure', 'Kebocoran informasi sensitif', 'Data pelanggan bocor'],
        ['Denial of Service', 'Membuat layanan tidak tersedia', 'Serangan DDoS'],
        ['Elevation of Privilege', 'Mendapatkan hak akses lebih tinggi', 'User biasa menjadi admin'],
    ]
)

# DREAD
doc.add_heading('b) DREAD (Risk Rating)', level=3)
doc.add_paragraph('DREAD adalah metode penilaian risiko untuk memprioritaskan ancaman. Setiap faktor dinilai skala 1-10.')
add_table(doc,
    ['Faktor', 'Pertanyaan Kunci'],
    [
        ['Damage', 'Seberapa besar kerusakan jika serangan berhasil?'],
        ['Reproducibility', 'Seberapa mudah serangan diulang?'],
        ['Exploitability', 'Seberapa mudah melakukan serangan?'],
        ['Affected Users', 'Berapa banyak pengguna terdampak?'],
        ['Discoverability', 'Seberapa mudah kerentanan ditemukan?'],
    ]
)
add_info_box(doc, "Total skor DREAD = D + R + E + A + D. Semakin tinggi skor, semakin prioritas untuk ditangani.", "PENTING")

# PASTA
doc.add_heading('c) PASTA (Process for Attack Simulation and Threat Analysis)', level=3)
doc.add_paragraph('PASTA memiliki 7 tahapan:')
pasta_steps = [
    'Definisi Objektif Bisnis',
    'Definisi Lingkup Teknis',
    'Dekomposisi Aplikasi',
    'Analisis Ancaman',
    'Analisis Kerentanan',
    'Enumerasi Serangan',
    'Analisis Risiko & Dampak'
]
for s in pasta_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('1.4 Langkah-Langkah Threat Modeling', level=2)
steps = [
    'Identifikasi Aset → Apa yang dilindungi?',
    'Buat Diagram Arsitektur → DFD (Data Flow Diagram)',
    'Identifikasi Ancaman → Gunakan STRIDE',
    'Identifikasi Kerentanan → Titik lemah sistem',
    'Evaluasi Risiko → Gunakan DREAD',
    'Tentukan Mitigasi → Solusi keamanan',
    'Validasi → Uji ulang'
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('1.5 Studi Kasus: Sistem E-Commerce "TokoOnline"', level=2)
doc.add_paragraph(
    'Skenario: Sebuah platform e-commerce "TokoOnline" memiliki fitur: login, katalog produk, keranjang belanja, pembayaran, dan riwayat pesanan.'
)
add_bold_paragraph(doc, 'Analisis STRIDE pada TokoOnline:')
add_table(doc,
    ['Ancaman', 'Contoh pada TokoOnline', 'Mitigasi'],
    [
        ['Spoofing', 'Attacker login dengan kredensial curian', 'MFA, CAPTCHA, rate limiting'],
        ['Tampering', 'Mengubah harga produk di request', 'Server-side validation, digital signature'],
        ['Repudiation', 'Pembeli menyangkal transaksi', 'Logging lengkap, audit trail'],
        ['Info Disclosure', 'SQL Injection membocorkan data kartu kredit', 'Parameterized queries, enkripsi data'],
        ['DoS', 'Flooding server dengan request', 'WAF, CDN, rate limiting'],
        ['EoP', 'Mengubah role dari buyer ke admin', 'RBAC, session validation'],
    ]
)

doc.add_heading('1.6 Data Flow Diagram (DFD) dalam Threat Modeling', level=2)
doc.add_paragraph('DFD digunakan untuk memetakan aliran data dalam sistem dan mengidentifikasi trust boundary (batas kepercayaan).')
add_bold_paragraph(doc, 'Elemen DFD:')
dfd_elements = [
    'External Entity (kotak) — pengguna/sistem eksternal',
    'Process (lingkaran) — proses yang mengolah data',
    'Data Store (garis paralel) — penyimpanan data',
    'Data Flow (panah) — arah aliran data',
    'Trust Boundary (garis putus-putus) — batas kepercayaan'
]
for e in dfd_elements:
    doc.add_paragraph(e, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# BAB 2: BRAINSTORMING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('BAB 2: BRAINSTORMING', level=1)

doc.add_heading('2.1 Definisi Brainstorming dalam Software Security', level=2)
doc.add_paragraph(
    'Brainstorming dalam konteks software security adalah teknik kolaboratif untuk mengidentifikasi potensi ancaman, kerentanan, dan skenario serangan '
    'dengan melibatkan berbagai perspektif dari tim pengembang, security analyst, dan stakeholder.'
)

doc.add_heading('2.2 Scenario Analysis (Analisis Skenario)', level=2)
doc.add_paragraph('Teknik untuk mengeksplorasi kemungkinan kejadian di masa depan dengan membuat skenario "bagaimana jika" (what-if) terkait keamanan sistem.')

add_bold_paragraph(doc, 'Langkah-langkah:')
sa_steps = [
    'Tentukan ruang lingkup — sistem/fitur yang dianalisis',
    'Identifikasi aktor — siapa saja yang berinteraksi (user, admin, attacker)',
    'Buat skenario — rangkaian kejadian yang mungkin terjadi',
    'Analisis dampak — apa konsekuensi dari setiap skenario',
    'Tentukan respons — bagaimana sistem harus merespons'
]
for s in sa_steps:
    doc.add_paragraph(s, style='List Number')

add_bold_paragraph(doc, 'Contoh: Aplikasi Mobile Banking')
add_table(doc,
    ['Skenario', 'Deskripsi', 'Dampak', 'Mitigasi'],
    [
        ['Skenario 1', 'User kehilangan HP yang sudah login', 'Akses tidak sah ke rekening', 'Auto-logout, biometric lock, remote wipe'],
        ['Skenario 2', 'MITM pada WiFi publik', 'Pencurian data transaksi', 'Certificate pinning, E2E encryption'],
        ['Skenario 3', 'Insider threat dari developer', 'Backdoor dalam kode', 'Code review, least privilege, audit'],
        ['Skenario 4', 'API key bocor di GitHub', 'Akses tidak sah ke backend', 'Secret management, key rotation'],
    ]
)

doc.add_heading('2.3 Premortem Analysis', level=2)
doc.add_paragraph(
    'Teknik di mana tim membayangkan bahwa proyek/sistem telah GAGAL dan kemudian bekerja mundur untuk mengidentifikasi penyebab kegagalan. '
    'Berbeda dengan postmortem yang dilakukan setelah kejadian, premortem dilakukan SEBELUM sistem di-deploy.'
)

add_bold_paragraph(doc, 'Prinsip Utama:')
pm_princ = [
    'Asumsikan sistem SUDAH diretas/gagal',
    'Tim diminta menjelaskan MENGAPA kegagalan terjadi',
    'Fokus pada PENYEBAB, bukan solusi (di tahap awal)'
]
for s in pm_princ:
    doc.add_paragraph(s, style='List Bullet')

add_bold_paragraph(doc, 'Langkah-langkah Premortem:')
pm_steps = [
    '"Sistem kita telah diretas dan data bocor. Ini adalah berita utama di media."',
    'Setiap anggota tim menulis kemungkinan penyebab (individual brainstorm)',
    'Semua penyebab dikumpulkan dan dibahas bersama (sharing)',
    'Penyebab diurutkan berdasarkan kemungkinan dan dampak (prioritasi)',
    'Buat rencana pencegahan untuk setiap penyebab (action plan)'
]
for s in pm_steps:
    doc.add_paragraph(s, style='List Number')

add_bold_paragraph(doc, 'Contoh Premortem — Sistem Rekam Medis Online (EHR):')
p = doc.add_paragraph()
r = p.add_run('Headline fiktif: "Data 10 Juta Pasien Bocor dari Sistem EHR — Termasuk Riwayat Penyakit dan NIK"')
r.italic = True
r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

add_table(doc,
    ['#', 'Penyebab Potensial', 'Kemungkinan', 'Dampak', 'Prioritas'],
    [
        ['1', 'SQL Injection pada form pencarian pasien', 'Tinggi', 'Kritis', 'P1'],
        ['2', 'Tidak ada enkripsi pada database backup', 'Tinggi', 'Kritis', 'P1'],
        ['3', 'Password admin default tidak diganti', 'Sedang', 'Kritis', 'P2'],
        ['4', 'API endpoint tidak terotentikasi', 'Tinggi', 'Tinggi', 'P1'],
        ['5', 'Log file berisi data sensitif tanpa proteksi', 'Sedang', 'Tinggi', 'P2'],
        ['6', 'Third-party library memiliki CVE belum dipatch', 'Tinggi', 'Tinggi', 'P1'],
    ]
)

doc.add_heading('2.4 Movie Plotting (Plot Film)', level=2)
doc.add_paragraph(
    'Teknik brainstorming kreatif di mana tim membuat skenario serangan seperti plot film thriller/action. '
    'Tim berperan sebagai "penjahat" dan merancang serangan paling dramatis dan merusak terhadap sistem.'
)

add_bold_paragraph(doc, 'Tujuan:')
mp_obj = [
    'Berpikir seperti attacker (adversarial thinking)',
    'Menemukan skenario serangan yang tidak terpikirkan dengan metode konvensional',
    'Mendorong kreativitas dan out-of-the-box thinking'
]
for s in mp_obj:
    doc.add_paragraph(s, style='List Bullet')

add_bold_paragraph(doc, 'Langkah-langkah Movie Plotting:')
mp_steps = [
    'Tentukan "film" — berikan judul dramatis untuk skenario',
    'Tentukan "tokoh" — siapa attacker, apa motivasinya',
    'Buat "plot" — rangkaian langkah serangan detail',
    'Tentukan "klimaks" — dampak terburuk yang terjadi',
    'Buat "twist" — bagaimana serangan bisa dicegah'
]
for s in mp_steps:
    doc.add_paragraph(s, style='List Number')

add_bold_paragraph(doc, 'Contoh Movie Plot: "The Silent Breach — Peretasan Tanpa Jejak"')
doc.add_paragraph('Villain: Mantan karyawan IT yang dipecat (insider threat)')
doc.add_paragraph('Target: Sistem perbankan digital "BankKu"')
doc.add_paragraph('Motivasi: Dendam dan keuntungan finansial')

add_bold_paragraph(doc, 'Plot:')
acts = [
    "ACT 1 - PERSIAPAN: Villain masih memiliki VPN credential yang belum di-revoke. Mempelajari arsitektur sistem yang ia bantu kembangkan. Membuat malware custom yang tidak terdeteksi antivirus.",
    "ACT 2 - INFILTRASI: Connect via VPN pada jam 2 pagi (low monitoring). Menggunakan credential lama untuk akses sistem internal. Memasang backdoor di server development. Lateral movement ke server production.",
    "ACT 3 - EKSEKUSI: Mengubah logic transfer untuk menyisipkan transaksi mikro. Setiap transaksi dibulatkan ke bawah, selisih masuk ke rekening villain. Total kerugian: Rp 50 Miliar dalam 6 bulan.",
    "ACT 4 - RESOLUSI (Pencegahan): Implementasi offboarding procedure yang ketat. Revoke semua credential saat karyawan keluar. Monitoring anomali transaksi dengan AI/ML. Zero Trust Architecture. Regular access review."
]
for a in acts:
    doc.add_paragraph(a, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# BAB 3: ATTACK TREE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('BAB 3: ATTACK TREE', level=1)

doc.add_heading('3.1 Definisi Attack Tree', level=2)
doc.add_paragraph(
    'Attack Tree adalah model hierarkis berbentuk pohon yang merepresentasikan berbagai cara (jalur serangan) '
    'untuk mencapai tujuan serangan tertentu. Dikembangkan oleh Bruce Schneier pada tahun 1999.'
)

doc.add_heading('3.2 Struktur Attack Tree', level=2)
doc.add_paragraph('Elemen-elemen Attack Tree:')
at_elements = [
    'Root Node (Akar): Tujuan utama attacker',
    'Intermediate Node: Sub-tujuan yang perlu dicapai',
    'Leaf Node (Daun): Aksi konkret/spesifik yang dilakukan attacker',
    'AND Node: SEMUA child harus tercapai agar parent goal berhasil',
    'OR Node: Cukup SALAH SATU child yang berhasil'
]
for e in at_elements:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('3.3 Notasi Attack Tree', level=2)
add_table(doc,
    ['Simbol', 'Arti'],
    [
        ['OR', 'Cukup salah satu jalur berhasil'],
        ['AND', 'Semua jalur harus berhasil secara bersamaan'],
        ['[P]', 'Possible — serangan mungkin dilakukan'],
        ['[I]', 'Impossible — serangan tidak mungkin/sangat sulit'],
        ['Cost', 'Biaya yang dibutuhkan attacker'],
        ['Probability', 'Kemungkinan keberhasilan serangan'],
    ]
)

doc.add_heading('3.4 Contoh Attack Tree: Mencuri Data Pengguna', level=2)
add_bold_paragraph(doc, 'ROOT: Mencuri Data Pengguna dari Aplikasi Web')
tree_lines = [
    "├── [OR] 1. Melalui Jaringan",
    "│   ├── [OR] 1.1 Man-in-the-Middle Attack",
    "│   │   ├── 1.1.1 ARP Spoofing [P, Cost: Rendah]",
    "│   │   └── 1.1.2 DNS Hijacking [P, Cost: Sedang]",
    "│   └── [OR] 1.2 Sniffing Traffic",
    "│       ├── 1.2.1 HTTP tanpa SSL [P, Cost: Rendah]",
    "│       └── 1.2.2 Downgrade Attack TLS [P, Cost: Tinggi]",
    "│",
    "├── [OR] 2. Melalui Aplikasi",
    "│   ├── [OR] 2.1 Injection Attack",
    "│   │   ├── 2.1.1 SQL Injection [P, Cost: Rendah]",
    "│   │   ├── 2.1.2 XSS mencuri session [P, Cost: Rendah]",
    "│   │   └── 2.1.3 LDAP Injection [P, Cost: Sedang]",
    "│   └── [AND] 2.2 Authentication Bypass",
    "│       ├── 2.2.1 Brute force password [P, Cost: Rendah]",
    "│       └── 2.2.2 Bypass MFA [P, Cost: Tinggi]",
    "│",
    "├── [OR] 3. Melalui Infrastruktur",
    "│   ├── 3.1 Server unpatched [P, Cost: Sedang]",
    "│   └── 3.2 Backup DB tidak terenkripsi [P, Cost: Rendah]",
    "│",
    "└── [OR] 4. Social Engineering",
    "    ├── 4.1 Phishing ke admin [P, Cost: Rendah]",
    "    └── 4.2 Pretexting ke helpdesk [P, Cost: Rendah]",
]
add_code_block(doc, tree_lines)

doc.add_heading('3.5 Kegunaan Attack Tree', level=2)
at_uses = [
    'Visualisasi jalur serangan secara komprehensif',
    'Prioritasi mitigasi berdasarkan cost/probability',
    'Komunikasi risiko ke stakeholder non-teknis',
    'Dokumentasi analisis keamanan',
    'Compliance — memenuhi standar keamanan (ISO 27001, NIST)'
]
for u in at_uses:
    doc.add_paragraph(u, style='List Number')

doc.add_heading('3.6 Cara Membuat Attack Tree', level=2)
at_steps = [
    'Tentukan root goal — apa tujuan attacker?',
    'Dekomposisi — pecah menjadi sub-goal',
    'Identifikasi leaf — tentukan aksi spesifik',
    'Tentukan relasi — AND atau OR',
    'Beri atribut — cost, probability, possible/impossible',
    'Analisis — temukan jalur termurah/termudah',
    'Mitigasi — prioritaskan pencegahan pada jalur berisiko tinggi'
]
for s in at_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# BAB 4: CHESS STRATEGY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('BAB 4: CHESS STRATEGY', level=1)

doc.add_heading('4.1 Konsep Chess Strategy dalam Software Security', level=2)
doc.add_paragraph(
    'Dalam software security, Chess Strategy adalah pendekatan yang menganalogikan keamanan perangkat lunak dengan permainan catur. '
    'Defenders (tim keamanan) dan Attackers (penyerang) saling berusaha mengalahkan satu sama lain.'
)

doc.add_heading('4.2 Analogi Catur-Security', level=2)
add_table(doc,
    ['Konsep Catur', 'Konsep Security', 'Penjelasan'],
    [
        ['Raja (King)', 'Aset paling berharga', 'Data sensitif, core system yang harus dilindungi'],
        ['Benteng (Rook)', 'Firewall, WAF', 'Pertahanan utama yang kuat dan langsung'],
        ['Kuda (Knight)', 'Penetration Tester', 'Bergerak tidak terduga, menguji pertahanan'],
        ['Gajah (Bishop)', 'IDS/IPS', 'Monitoring jarak jauh/diagonal'],
        ['Menteri (Queen)', 'SIEM', 'Paling powerful, multi-fungsi'],
        ['Pion (Pawn)', 'Defense in Depth', 'Lapisan pertahanan pertama/terluar'],
        ['Skak Mat', 'Full System Compromise', 'Attacker berhasil menguasai seluruh sistem'],
        ['Gambit', 'Honeypot', 'Mengorbankan aset palsu untuk menjebak attacker'],
    ]
)

doc.add_heading('4.3 Prinsip-Prinsip Chess Strategy', level=2)

principles = [
    ("a) Think Ahead (Berpikir ke Depan)",
     "Catur: Pemain harus berpikir beberapa langkah ke depan.",
     "Security: Proactive security — antisipasi serangan sebelum terjadi.",
     "Implementasi: Threat modeling, security by design, SAST/DAST di CI/CD."),
    ("b) Control the Center (Kontrol Pusat)",
     "Catur: Menguasai pusat papan memberikan keuntungan strategis.",
     "Security: Kontrol aset kritis — lindungi core system terlebih dahulu.",
     "Implementasi: Zero Trust Architecture, identity management."),
    ("c) Protect the King (Lindungi Raja)",
     "Catur: Raja harus dilindungi dengan segala cara.",
     "Security: Lindungi crown jewels — data dan sistem paling kritis.",
     "Implementasi: Enkripsi, access control, segmentasi jaringan."),
    ("d) Defense in Depth (Pertahanan Berlapis)",
     "Catur: Tidak mengandalkan satu bidak untuk pertahanan.",
     "Security: Multiple layers of security.",
     "Implementasi: Firewall + IDS + WAF + Enkripsi + MFA."),
    ("e) Sacrifice / Gambit (Pengorbanan Strategis)",
     "Catur: Mengorbankan bidak untuk keuntungan posisi.",
     "Security: Honeypot/Honeynet — menjebak attacker dengan aset palsu.",
     "Implementasi: Decoy systems, canary tokens."),
    ("f) Tempo & Initiative",
     "Catur: Mengambil inisiatif dan tidak bermain reaktif.",
     "Security: Incident response yang cepat, patching proaktif.",
     "Implementasi: SOAR (Security Orchestration, Automation and Response)."),
]
for title, catur, sec, impl in principles:
    add_bold_paragraph(doc, title, 12)
    doc.add_paragraph(f"• {catur}")
    doc.add_paragraph(f"• {sec}")
    doc.add_paragraph(f"• {impl}")

doc.add_heading('4.4 Penerapan Chess Strategy — FinTech', level=2)
add_bold_paragraph(doc, 'OPENING (Fase Awal — Design & Development):')
for s in ['Threat Modeling sejak awal (STRIDE/DREAD)', 'Secure coding standards', 'Security requirements dalam backlog', 'SAST (Static Analysis) di CI/CD pipeline']:
    doc.add_paragraph(s, style='List Bullet')

add_bold_paragraph(doc, 'MIDDLE GAME (Fase Operasional):')
for s in ['Monitoring real-time (SIEM)', 'Vulnerability scanning berkala', 'Penetration testing', 'Security awareness training', 'Incident response drill']:
    doc.add_paragraph(s, style='List Bullet')

add_bold_paragraph(doc, 'END GAME (Fase Respons Insiden):')
for s in ['Incident detection & triage', 'Containment & eradication', 'Recovery & restoration', 'Post-incident review (lessons learned)', 'Update defense strategy']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# BAB 5: SOAL DAN JAWABAN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('BAB 5: SOAL DAN JAWABAN', level=1)

# ── PG ──
doc.add_heading('A. Soal Pilihan Ganda', level=2)

pg_soal = [
    ('1. Dalam metodologi STRIDE, ancaman "Tampering" mengacu pada:',
     ['A. Menyamar sebagai pengguna lain', 'B. Mengubah data tanpa izin', 'C. Membuat layanan tidak tersedia', 'D. Mendapatkan hak akses lebih tinggi'],
     'Jawaban: B. Mengubah data tanpa izin — Tampering berarti mengubah/memodifikasi data tanpa otorisasi.'),
    ('2. Faktor mana dalam DREAD yang mengevaluasi "seberapa mudah serangan diulangi"?',
     ['A. Damage', 'B. Reproducibility', 'C. Exploitability', 'D. Discoverability'],
     'Jawaban: B. Reproducibility — Faktor R dalam DREAD mengukur kemudahan mereproduksi serangan.'),
    ('3. Trust boundary dalam DFD ditandai dengan:',
     ['A. Garis lurus', 'B. Lingkaran', 'C. Garis putus-putus', 'D. Kotak'],
     'Jawaban: C. Garis putus-putus — Trust boundary menandai batas kepercayaan antar komponen.'),
    ('4. Metodologi PASTA memiliki berapa tahapan?',
     ['A. 5', 'B. 6', 'C. 7', 'D. 8'],
     'Jawaban: C. 7 — PASTA memiliki 7 tahapan dari definisi objektif bisnis hingga analisis risiko.'),
    ('5. Dalam STRIDE, "Repudiation" berarti:',
     ['A. Mencuri informasi', 'B. Menyangkal tindakan yang telah dilakukan', 'C. Menyerang ketersediaan layanan', 'D. Memodifikasi data'],
     'Jawaban: B. Menyangkal tindakan yang telah dilakukan.'),
    ('6. Siapa yang mengembangkan konsep Attack Tree?',
     ['A. Bill Gates', 'B. Bruce Schneier', 'C. Linus Torvalds', 'D. Kevin Mitnick'],
     'Jawaban: B. Bruce Schneier — Mengembangkan Attack Tree pada tahun 1999.'),
    ('7. Dalam Attack Tree, AND node berarti:',
     ['A. Cukup satu jalur berhasil', 'B. Semua jalur harus berhasil', 'C. Tidak ada jalur yang perlu berhasil', 'D. Jalur opsional'],
     'Jawaban: B. Semua jalur harus berhasil — AND node mengharuskan semua child tercapai.'),
    ('8. Konsep "Gambit" dalam Chess Strategy dianalogikan dengan:',
     ['A. Firewall', 'B. Honeypot', 'C. Enkripsi', 'D. Backup'],
     'Jawaban: B. Honeypot — Mengorbankan aset palsu untuk menjebak attacker.'),
    ('9. Apa perbedaan utama Premortem dengan Postmortem?',
     ['A. Premortem dilakukan setelah insiden', 'B. Premortem dilakukan sebelum insiden', 'C. Keduanya sama', 'D. Postmortem dilakukan sebelum insiden'],
     'Jawaban: B. Premortem dilakukan sebelum insiden — Tim membayangkan kegagalan yang belum terjadi.'),
    ('10. Dalam analogi Chess Strategy, "Raja" merepresentasikan:',
     ['A. Firewall', 'B. Aset paling berharga', 'C. Penetration tester', 'D. IDS/IPS'],
     'Jawaban: B. Aset paling berharga — Raja = crown jewels = data/sistem paling kritis.'),
]

for soal_text, options, jawaban in pg_soal:
    add_bold_paragraph(doc, soal_text)
    for opt in options:
        doc.add_paragraph(opt, style='List Bullet')
    p = doc.add_paragraph()
    r = p.add_run(jawaban)
    r.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    r.bold = True
    doc.add_paragraph()

# ── Essay ──
doc.add_heading('B. Soal Essay', level=2)

# Essay 1
add_bold_paragraph(doc, 'Soal 1: Jelaskan perbedaan antara Premortem dan Postmortem dalam konteks software security!', 12)
doc.add_paragraph(
    'Jawaban:\n'
    '• Premortem dilakukan SEBELUM sistem di-deploy atau insiden terjadi. Tim membayangkan bahwa sistem telah gagal/diretas dan bekerja mundur untuk mengidentifikasi penyebab potensial. Tujuannya preventif.\n'
    '• Postmortem dilakukan SETELAH insiden terjadi. Tim menganalisis apa yang terjadi, mengapa terjadi, dan bagaimana mencegah kejadian serupa. Tujuannya korektif.\n'
    '• Keuntungan premortem: Lebih murah dan efektif karena masalah dicegah sebelum terjadi. Menghilangkan bias "ini tidak akan terjadi" karena tim diminta mengasumsikan kegagalan sudah terjadi.'
)

# Essay 2
add_bold_paragraph(doc, 'Soal 2: Buatlah contoh Movie Plot untuk serangan terhadap sistem Smart Home!', 12)
doc.add_paragraph(
    'Jawaban:\n'
    'Judul: "Home Invasion 2.0"\n'
    'Villain: Hacker yang menarget keluarga kaya melalui perangkat IoT\n\n'
    'Plot:\n'
    '1. Attacker menemukan smart lock merk X memiliki vulnerability Bluetooth (CVE publik)\n'
    '2. Attacker melakukan wardriving di area perumahan mewah\n'
    '3. Menemukan target dengan smart lock vulnerable\n'
    '4. Melakukan exploit Bluetooth untuk membuka kunci pintu\n'
    '5. Mematikan CCTV melalui smart home hub yang sudah terkompromi\n'
    '6. Mencuri barang berharga tanpa jejak digital\n\n'
    'Pencegahan: Firmware update otomatis, network segmentation IoT, anomaly detection, physical backup lock.'
)

# Essay 3
add_bold_paragraph(doc, 'Soal 3: Gambarkan attack tree untuk "Mengambil alih akun admin WordPress"!', 12)
doc.add_paragraph(
    'Jawaban:\n'
    'ROOT: Mengambil Alih Akun Admin WordPress\n\n'
    '[OR] 1. Credential Attack\n'
    '  • 1.1 Brute Force wp-login.php [P, Cost: Rendah]\n'
    '  • 1.2 Credential Stuffing dari data breach [P, Cost: Rendah]\n'
    '  • 1.3 Phishing ke admin [P, Cost: Rendah]\n\n'
    '[OR] 2. Exploit Vulnerability\n'
    '  • 2.1 Plugin vulnerability (RCE) [P, Cost: Sedang]\n'
    '  • 2.2 Theme vulnerability (LFI/RFI) [P, Cost: Sedang]\n'
    '  • 2.3 WordPress core exploit (0-day) [P, Cost: Tinggi]\n\n'
    '[OR] 3. Session Hijacking\n'
    '  • [AND] 3.1 XSS + Cookie Stealing (harus dua-duanya berhasil)\n'
    '  • 3.2 Session Fixation [P, Cost: Sedang]\n\n'
    '[OR] 4. Server-Level Attack\n'
    '  • 4.1 Akses database langsung (MySQL) [P, Cost: Tinggi]\n'
    '  • 4.2 File upload vulnerability → webshell [P, Cost: Sedang]\n\n'
    'Analisis: Jalur termurah = Brute Force dan Credential Stuffing. Prioritas mitigasi: rate limiting, MFA, strong password.'
)

# Essay 4
add_bold_paragraph(doc, 'Soal 4: Jelaskan konsep "Defense in Depth" menggunakan analogi catur!', 12)
doc.add_paragraph(
    'Jawaban:\n'
    'Dalam catur, pemain yang baik tidak mengandalkan satu bidak untuk melindungi raja. Ia menggunakan formasi berlapis:\n\n'
    '• Pion (Layer 1): Firewall, input validation — pertahanan terluar\n'
    '• Knight/Bishop (Layer 2): IDS/IPS, WAF — deteksi dan pencegahan\n'
    '• Rook (Layer 3): Enkripsi data, access control — proteksi data\n'
    '• Queen/SIEM (Layer 4): Monitoring, logging — deteksi anomali\n'
    '• King Castle (Layer 5): Backup, disaster recovery — pertahanan terakhir\n\n'
    'Jika satu lapisan ditembus, lapisan berikutnya masih melindungi raja (aset kritis). '
    'Prinsipnya: jangan pernah mengandalkan satu mekanisme keamanan saja.'
)

# Essay 5
add_bold_paragraph(doc, 'Soal 5 (Integratif): Sebuah startup fintech akan meluncurkan aplikasi pembayaran digital. Lakukan analisis keamanan menggunakan semua 4 konsep!', 12)
doc.add_paragraph(
    'Jawaban:\n\n'
    'A. Threat Modeling (STRIDE):\n'
    '  • Spoofing: Pemalsuan identitas pembayar → Mitigasi: eKYC, biometric\n'
    '  • Tampering: Pengubahan jumlah transaksi → Mitigasi: digital signature\n'
    '  • Repudiation: Penyangkalan transaksi → Mitigasi: blockchain audit trail\n'
    '  • Info Disclosure: Kebocoran data keuangan → Mitigasi: encryption at rest & transit\n'
    '  • DoS: Serangan ke payment gateway → Mitigasi: rate limiting, redundancy\n'
    '  • EoP: Eskalasi ke admin dashboard → Mitigasi: strict RBAC, MFA admin\n\n'
    'B. Brainstorming - Premortem:\n'
    '  Headline: "Jutaan Rupiah Raib — Aplikasi FinTech ABC Diretas Hacker"\n'
    '  • Penyebab: API payment tidak tervalidasi properly\n'
    '  • Penyebab: Hardcoded API key di mobile app\n'
    '  • Penyebab: Tidak ada fraud detection system\n\n'
    'C. Attack Tree:\n'
    '  ROOT: Mencuri Uang dari Aplikasi Pembayaran\n'
    '  [OR] API Exploitation → Man-in-the-middle\n'
    '  [OR] Reverse Engineering APK → Extract API keys\n'
    '  [OR] Social Engineering → Phishing ke customer\n'
    '  [AND] Account Takeover → Credential + Bypass OTP\n\n'
    'D. Chess Strategy:\n'
    '  • Opening: Security by design, threat modeling\n'
    '  • Middle: Fraud detection AI, real-time monitoring\n'
    '  • Endgame: Incident response plan, cyber insurance\n'
    '  • Gambit: Honeypot account untuk deteksi fraud'
)

doc.add_page_break()

# ── RANGKUMAN ──
doc.add_heading('RANGKUMAN KUNCI', level=1)
add_table(doc,
    ['Topik', 'Kata Kunci', 'Poin Penting'],
    [
        ['Threat Modeling', 'STRIDE, DREAD, DFD, Trust Boundary', 'Proses sistematis identifikasi ancaman'],
        ['Scenario Analysis', 'What-if, Aktor, Dampak, Respons', 'Eksplorasi kemungkinan serangan'],
        ['Premortem', '"Sudah gagal", Kerja mundur, Preventif', 'Imajinasi kegagalan sebelum terjadi'],
        ['Movie Plotting', 'Adversarial thinking, Plot, Kreatif', 'Berpikir seperti attacker'],
        ['Attack Tree', 'Root, AND/OR, Leaf, Cost, Probability', 'Visualisasi hierarkis jalur serangan'],
        ['Chess Strategy', 'Defense in Depth, Think Ahead, Gambit', 'Analogi strategis keamanan'],
    ]
)

doc.add_paragraph()
add_bold_paragraph(doc, 'Tips UTS:', 14)
tips = [
    'Pahami STRIDE dan bisa mengaplikasikan ke studi kasus apapun',
    'Bisa membuat Attack Tree lengkap dengan notasi AND/OR',
    'Pahami perbedaan Premortem vs Postmortem',
    'Hafal analogi Chess Strategy dan penerapannya',
    'Latih membuat skenario brainstorming untuk berbagai sistem'
]
for t in tips:
    doc.add_paragraph(t, style='List Bullet')

# ── Save ──
output_path = r'c:\SEMESTER4\SOFTSEC\Materi_UTS_Software_Security_BINUS.docx'
doc.save(output_path)
print(f"✅ Word document saved to: {output_path}")
