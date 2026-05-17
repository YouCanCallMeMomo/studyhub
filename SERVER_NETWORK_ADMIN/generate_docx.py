from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# STYLES & FORMATTING
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Colors
DARK_BLUE = RGBColor(0x0D, 0x1B, 0x2A)
NAVY = RGBColor(0x1B, 0x26, 0x3B)
ACCENT_BLUE = RGBColor(0x0F, 0x4C, 0x81)
ACCENT_TEAL = RGBColor(0x00, 0x7B, 0x83)
ACCENT_RED = RGBColor(0xC0, 0x39, 0x2B)
ACCENT_GREEN = RGBColor(0x27, 0xAE, 0x60)
ACCENT_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF0, 0xF0, 0xF5)
MEDIUM_GRAY = RGBColor(0x5D, 0x6D, 0x7E)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_heading(text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def add_body(text, bold=False, italic=False, color=None, size=Pt(11), align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_bullet(text, bold_prefix="", level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = color
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_table(headers, rows, header_color="0D1B2A", widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F0F5")
    
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    doc.add_paragraph()
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    # light bg via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_info_box(title, text, color_hex="E8F4FD", border_color=ACCENT_BLUE):
    p = doc.add_paragraph()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {title}: ")
    run.bold = True
    run.font.color.rgb = border_color
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    return p

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 80)
    run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
    run.font.size = Pt(8)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UAS — SERVER NETWORK ADMINISTRATION")
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("DevSecOps Pipeline Security")
run.font.size = Pt(20)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("━" * 40)
run.font.color.rgb = ACCENT_TEAL
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Study Case & Materi Komprehensif")
run.font.size = Pt(16)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'
run.italic = True

for _ in range(3):
    doc.add_paragraph()

info_lines = [
    ("Universitas", "BINUS University"),
    ("Program Studi", "Cyber Security"),
    ("Semester", "4 (Empat)"),
    ("Jenis Ujian", "Ujian Akhir Semester (UAS)"),
    ("Mata Kuliah", "Server Network Administration"),
    ("Topik", "DevSecOps — CI/CD Pipeline Security"),
    ("Tanggal", "April 2026"),
]

table = doc.add_table(rows=len(info_lines), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate(info_lines):
    cell0 = table.rows[i].cells[0]
    cell1 = table.rows[i].cells[1]
    cell0.width = Cm(4)
    cell1.width = Cm(8)
    p0 = cell0.paragraphs[0]
    run0 = p0.add_run(label)
    run0.bold = True
    run0.font.name = 'Calibri'
    run0.font.size = Pt(11)
    run0.font.color.rgb = ACCENT_BLUE
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run(f"  {value}")
    run1.font.name = 'Calibri'
    run1.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (Manual)
# ============================================================
add_styled_heading("DAFTAR ISI", level=1)
add_separator()

toc_items = [
    ("BAGIAN A", "Penjelasan Materi DevSecOps"),
    ("  A.1", "DevOps vs DevSecOps"),
    ("  A.2", "Tahapan Pipeline DevSecOps"),
    ("  A.3", "Tools DevSecOps yang Wajib Diketahui"),
    ("", ""),
    ("BAGIAN B", "Study Case — Soal Ujian"),
    ("  ", "Latar Belakang: PT. NusaCloud Technologies"),
    ("  ", "Insiden Keamanan"),
    ("  Soal 1", "Analisis Risiko Pipeline (20 Poin)"),
    ("  Soal 2", "Perbaikan yang Dapat Dilakukan (20 Poin)"),
    ("  Soal 3", "Redesign Pipeline (20 Poin)"),
    ("  Soal 4", "What-If Scenario (20 Poin)"),
    ("  Soal 5", "Critical Missing Process (20 Poin)"),
    ("", ""),
    ("BAGIAN C", "Quick Reference — Tools DevSecOps"),
    ("BAGIAN D", "Tips Menjawab Ujian"),
]

for prefix, item in toc_items:
    if not prefix and not item:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    if prefix.strip():
        run = p.add_run(f"{prefix}  ")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT_BLUE
    run = p.add_run(item)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# BAGIAN A: PENJELASAN MATERI DEVSECOPS
# ============================================================
add_styled_heading("BAGIAN A: PENJELASAN MATERI DEVSECOPS", level=1, color=DARK_BLUE)
add_separator()

# A.1 Apa Itu DevSecOps
add_styled_heading("A.1 — Apa Itu DevSecOps?", level=2, color=ACCENT_BLUE)

add_body("DevSecOps adalah evolusi dari DevOps yang mengintegrasikan Security (Sec) ke dalam setiap tahap siklus pengembangan perangkat lunak. Filosofi utamanya adalah \"Shift Left\" — menggeser aktivitas keamanan ke tahap paling awal dalam pipeline, bukan menunggu hingga akhir deployment.")

add_info_box("PENTING", "Dalam DevSecOps, keamanan adalah tanggung jawab SEMUA orang dalam tim (Dev + Sec + Ops), bukan hanya tim security.", "FFF3CD", ACCENT_ORANGE)

add_styled_heading("Perbandingan: DevOps vs DevSecOps", level=3, color=NAVY)

add_table(
    ["Aspek", "DevOps", "DevSecOps"],
    [
        ["Fokus", "Speed & Automation", "Speed, Automation & Security"],
        ["Security", "Ditambahkan di akhir (afterthought)", "Terintegrasi di setiap tahap"],
        ["Testing", "Functional & Performance", "Functional, Performance & Security"],
        ["Tanggung Jawab", "Dev + Ops", "Dev + Sec + Ops"],
        ["Pipeline", "CI/CD", "CI/CD dengan Security Gates"],
        ["Kultur", "Kolaborasi Dev & Ops", "Kolaborasi Dev, Sec & Ops"],
    ]
)

# A.2 Tahapan Pipeline
add_styled_heading("A.2 — Tahapan Pipeline DevSecOps", level=2, color=ACCENT_BLUE)

add_body("Pipeline DevSecOps terdiri dari 6 tahap utama dengan security terintegrasi di setiap level:")

add_table(
    ["Stage", "Aktivitas", "Security Integration"],
    [
        ["1. PLAN", "Perencanaan fitur & arsitektur", "Threat Modeling, Risk Assessment"],
        ["2. CODE", "Penulisan kode", "IDE Security Plugins, Pre-commit Hooks, Secret Detection"],
        ["3. BUILD", "Kompilasi & packaging", "SAST, SCA, License Compliance Check"],
        ["4. TEST", "Pengujian", "DAST, IAST, Penetration Testing, Fuzzing"],
        ["5. DEPLOY", "Deployment ke server", "Container Scanning, IaC Scanning, Config Audit"],
        ["6. MONITOR", "Monitoring production", "SIEM, IDS/IPS, Runtime Protection, Alerting"],
    ]
)

# A.3 Tools DevSecOps
add_styled_heading("A.3 — Tools DevSecOps yang Wajib Diketahui", level=2, color=ACCENT_BLUE)

# SAST
add_styled_heading("1. SAST (Static Application Security Testing)", level=3, color=NAVY)
add_body("Menganalisis source code tanpa menjalankan aplikasi untuk menemukan kerentanan keamanan.", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["SonarQube", "Code quality & security analysis platform", "Community / Enterprise"],
        ["Semgrep", "Lightweight, customizable, sangat cepat", "Open Source"],
        ["Checkmarx", "Enterprise-grade SAST, deep data-flow analysis", "Commercial"],
        ["Snyk Code", "Developer-friendly, real-time feedback di IDE", "Freemium"],
        ["CodeQL (GitHub)", "Native integration untuk pengguna GitHub", "Free for OSS"],
    ]
)

add_body("Contoh Temuan SAST:", bold=True)
add_code_block("# ❌ VULNERABLE - SQL Injection\nquery = \"SELECT * FROM users WHERE id = \" + user_input\ncursor.execute(query)")
add_code_block("# ✅ SECURE - Parameterized Query\nquery = \"SELECT * FROM users WHERE id = %s\"\ncursor.execute(query, (user_input,))")

# SCA
add_styled_heading("2. SCA (Software Composition Analysis)", level=3, color=NAVY)
add_body("Mengidentifikasi kerentanan dalam library open-source dan dependencies.", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["Snyk", "Leader SCA, developer-first approach", "Freemium"],
        ["Trivy", "Lightweight, multi-purpose (SCA + Container + IaC)", "Open Source"],
        ["OWASP Dependency-Check", "Scanning dependencies terhadap CVE database", "Open Source"],
        ["Dependabot (GitHub)", "Auto-generate PR untuk update dependencies", "Free"],
    ]
)

# DAST
add_styled_heading("3. DAST (Dynamic Application Security Testing)", level=3, color=NAVY)
add_body("Menguji aplikasi yang sedang berjalan dengan mensimulasikan serangan dari luar.", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["OWASP ZAP", "Industry-standard open-source DAST", "Open Source"],
        ["Burp Suite", "Professional-grade web security testing", "Community / Pro"],
        ["Nikto", "Web server vulnerability scanner", "Open Source"],
        ["Nuclei", "Template-based vulnerability scanner", "Open Source"],
    ]
)

# Container Scanning
add_styled_heading("4. Container & Image Scanning", level=3, color=NAVY)
add_body("Mendeteksi kerentanan dalam container images (Docker, Kubernetes).", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["Trivy", "All-in-one scanner (container, IaC, SCA)", "Open Source"],
        ["Aqua Security", "Enterprise container security platform", "Commercial"],
        ["Grype", "Container image vulnerability scanner", "Open Source"],
        ["Docker Scout", "Native Docker security scanning", "Freemium"],
    ]
)

# IaC Scanning
add_styled_heading("5. Infrastructure as Code (IaC) Scanning", level=3, color=NAVY)
add_body("Memindai konfigurasi infrastruktur (Terraform, K8s, CloudFormation) untuk menemukan miskonfigurasi.", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["Checkov", "Policy-as-code untuk Terraform, K8s, CloudFormation", "Open Source"],
        ["Terrascan", "IaC compliance & security scanner", "Open Source"],
        ["tfsec", "Terraform-specific security scanner", "Open Source"],
        ["KICS", "Keeping Infrastructure as Code Secure", "Open Source"],
    ]
)

# Secret Management
add_styled_heading("6. Secret Management & Detection", level=3, color=NAVY)
add_body("Mendeteksi dan mengelola secrets (API keys, passwords, tokens) agar tidak bocor.", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["GitLeaks", "Scanning repo untuk hardcoded secrets", "Open Source"],
        ["TruffleHog", "Deteksi secrets di Git history", "Open Source"],
        ["HashiCorp Vault", "Enterprise secret management", "Open Source / Enterprise"],
        ["AWS Secrets Manager", "Cloud-native secret management", "Pay-per-use"],
    ]
)

add_body("Contoh Secret Leak:", bold=True)
add_code_block('# ❌ BERBAHAYA - Hardcoded\nAPI_KEY = "sk-proj-abc123def456ghi789"\nDATABASE_PASSWORD = "SuperSecret123!"')
add_code_block('# ✅ AMAN - Environment Variables\nimport os\nAPI_KEY = os.environ.get("API_KEY")\nDATABASE_PASSWORD = os.environ.get("DB_PASSWORD")')

# Monitoring
add_styled_heading("7. Monitoring & Runtime Protection", level=3, color=NAVY)
add_body("Memantau aplikasi dan infrastruktur secara real-time saat berjalan di production.", italic=True, color=MEDIUM_GRAY)

add_table(
    ["Tool", "Deskripsi", "Lisensi"],
    [
        ["Splunk", "Enterprise SIEM & log management", "Commercial"],
        ["ELK Stack", "Elasticsearch, Logstash, Kibana", "Open Source"],
        ["Prometheus + Grafana", "Metrics monitoring & visualization", "Open Source"],
        ["Falco", "Runtime security for Kubernetes", "Open Source"],
        ["Wazuh", "Open-source SIEM & XDR", "Open Source"],
    ]
)

doc.add_page_break()

# ============================================================
# BAGIAN B: STUDY CASE
# ============================================================
add_styled_heading("BAGIAN B: STUDY CASE — SOAL UJIAN", level=1, color=DARK_BLUE)
add_separator()

add_styled_heading("Insiden Keamanan di PT. NusaCloud Technologies", level=2, color=ACCENT_BLUE)

# Latar Belakang
add_styled_heading("Latar Belakang Perusahaan", level=3, color=NAVY)

add_body("PT. NusaCloud Technologies adalah perusahaan teknologi di Indonesia yang menyediakan platform cloud-based Enterprise Resource Planning (ERP) untuk lebih dari 500 perusahaan di Asia Tenggara. Sistem ERP mereka menangani data sensitif termasuk:")

add_bullet("Data keuangan perusahaan klien")
add_bullet("Informasi karyawan (gaji, NIK, alamat)")
add_bullet("Data transaksi bisnis")
add_bullet("Laporan pajak dan audit")

# Pipeline Original
add_styled_heading("Arsitektur Pipeline CI/CD Original (SEBELUM Insiden)", level=3, color=NAVY)

add_body("Pipeline CI/CD yang digunakan sebelum insiden sangat sederhana dan TIDAK memiliki security controls:")

add_table(
    ["Stage", "Komponen", "Masalah Keamanan"],
    [
        ["1. Code", "Developer push langsung ke branch main", "Tidak ada pull request atau code review"],
        ["2. Build", "Jenkins build otomatis", "Tidak ada SAST atau SCA scanning"],
        ["3. Registry", "Docker image push ke private registry", "Tidak ada vulnerability scanning"],
        ["4. Deploy", "Auto deploy ke production (AWS EC2)", "Tidak ada DAST testing"],
        ["5. Monitor", "Basic uptime monitoring (Ping)", "Tidak ada SIEM atau security monitoring"],
    ],
    header_color="C0392B"
)

add_info_box("⚠️ KELEMAHAN KRITIS", "Jenkins menggunakan shared admin credentials (15 developer); API keys hardcoded dalam source code; Tidak ada branch protection; Semua developer punya root access ke production; Tidak ada staging environment", "FDEDEC", ACCENT_RED)

# Insiden
add_styled_heading("INSIDEN YANG TERJADI", level=3, color=ACCENT_RED)

add_body("Pada 15 Maret 2026, tim security menerima laporan bahwa data keuangan klien muncul di dark web. Investigasi mengungkap:", bold=True)

add_table(
    ["Tanggal", "Kejadian"],
    [
        ["1 Jan 2026", "Mantan developer (resign Nov 2025) masih memiliki akses ke GitHub & Jenkins — tidak ada revocation"],
        ["10 Jan 2026", "Mantan developer login ke Jenkins menggunakan shared admin credentials yang masih valid"],
        ["15 Jan 2026", "Attacker memodifikasi Jenkinsfile, menambahkan backdoor (reverse shell) ke Docker image"],
        ["16 Jan – 14 Mar 2026", "Setiap build otomatis membawa backdoor ke production. Data di-exfiltrate via DNS tunneling"],
        ["15 Mar 2026", "Data klien muncul di dark web — insiden terdeteksi"],
    ],
    header_color="C0392B"
)

add_styled_heading("Dampak Insiden:", level=3, color=ACCENT_RED)
add_bullet("380 perusahaan ", bold_prefix="", color=ACCENT_RED)
p = doc.paragraphs[-1]
p.clear()
run = p.add_run("380 perusahaan ")
run.bold = True
run.font.name = 'Calibri'
run2 = p.add_run("klien terdampak kebocoran data")
run2.font.name = 'Calibri'

add_bullet("2.5 juta record data karyawan dan keuangan bocor")
add_bullet("Kerugian finansial diperkirakan Rp 45 Miliar")
add_bullet("120 klien memutuskan kontrak")
add_bullet("Sanksi dari Kementerian Kominfo (UU PDP)")

add_styled_heading("Root Cause Analysis:", level=3, color=NAVY)

add_table(
    ["#", "Root Cause", "Deskripsi"],
    [
        ["1", "No Access Revocation", "Mantan karyawan masih memiliki akses penuh"],
        ["2", "Shared Credentials", "Tidak ada accountability individual"],
        ["3", "No Pipeline Integrity Check", "Perubahan Jenkinsfile tidak diaudit"],
        ["4", "No Security Scanning", "Backdoor dalam Docker image tidak terdeteksi"],
        ["5", "No Security Monitoring", "DNS tunneling tidak terdeteksi selama 2 bulan"],
        ["6", "No Code Review", "Modifikasi pipeline tidak memerlukan approval"],
    ],
    header_color="C0392B"
)

doc.add_page_break()

# ============================================================
# SOAL 1
# ============================================================
add_styled_heading("SOAL 1: Analisis Risiko Pipeline yang Tidak Berubah (20 Poin)", level=2, color=ACCENT_RED)

add_info_box("PERTANYAAN", "Berdasarkan insiden yang terjadi, jelaskan mengapa sangat berbahaya jika pipeline CI/CD tetap digunakan tanpa perubahan. Identifikasi minimal 5 risiko kritikal dan jelaskan bagaimana masing-masing dapat dieksploitasi.", "E8F4FD", ACCENT_BLUE)

add_styled_heading("JAWABAN SOAL 1", level=3, color=ACCENT_GREEN)

# Risiko 1
add_body("Risiko 1: Supply Chain Attack via Build Pipeline (CRITICAL)", bold=True, color=ACCENT_RED, size=Pt(12))
add_body("Pipeline saat ini tidak memiliki integrity verification pada build process. Jenkinsfile dapat dimodifikasi oleh siapa saja tanpa approval atau audit trail.")
add_body("Bagaimana Dieksploitasi:", bold=True)
add_bullet("Attacker dapat memodifikasi build script untuk menyisipkan malicious code")
add_bullet("Mirip dengan serangan SolarWinds (2020) — malware SUNBURST disuntikkan ke build process")
add_bullet("Dalam SolarWinds, malware SUNSPOT memantau MSBuild dan mengganti source code saat kompilasi, lalu mengembalikan ke versi asli")
add_bullet("Binary ter-sign secara legitimate sehingga deteksi sangat sulit")
add_body("Dampak: Setiap deployment membawa malware; semua klien terinfeksi; deteksi bisa memakan waktu berbulan-bulan.", bold=True, color=ACCENT_RED)

doc.add_paragraph()

# Risiko 2
add_body("Risiko 2: Insider Threat Tanpa Accountability (HIGH)", bold=True, color=ACCENT_RED, size=Pt(12))
add_body("Shared admin credentials berarti tidak ada cara mengetahui SIAPA melakukan tindakan tertentu. Melanggar prinsip Non-Repudiation dan Least Privilege.")
add_body("Bagaimana Dieksploitasi:", bold=True)
add_bullet("15 developer bisa melakukan perubahan malicious tanpa dilacak")
add_bullet("Mantan karyawan dengan credentials bisa login kapan saja")
add_bullet("Semua aktivitas tercatat sebagai 'admin' — investigasi forensik mustahil")
add_bullet("Prinsip Zero Trust dilanggar total")

doc.add_paragraph()

# Risiko 3
add_body("Risiko 3: Credential Exposure via Hardcoded Secrets (HIGH)", bold=True, color=ACCENT_RED, size=Pt(12))
add_body("API keys dan database passwords yang di-hardcode menciptakan single point of compromise.")
add_body("Bagaimana Dieksploitasi:", bold=True)
add_bullet("Attacker yang akses repo langsung bisa akses database production")
add_bullet("Tools GitLeaks/TruffleHog otomatis mendeteksi secrets di repositories")
add_bullet("Bahkan jika dihapus dari kode, masih ada di Git history")
add_bullet("Rotasi credentials sangat sulit karena hardcoded di banyak tempat")

doc.add_paragraph()

# Risiko 4
add_body("Risiko 4: Undetected Vulnerabilities in Dependencies (HIGH)", bold=True, color=ACCENT_RED, size=Pt(12))
add_body("Tanpa SCA scanning dan container scanning, aplikasi bisa menggunakan library dengan CVE yang sudah diketahui.")
add_body("Bagaimana Dieksploitasi:", bold=True)
add_bullet("Contoh: Log4Shell (CVE-2021-44228) — Remote Code Execution via Log4j")
add_bullet("Attacker scan target, deteksi versi library vulnerable, gunakan exploit publik")
add_bullet("Docker image dengan OS packages outdated juga rentan")

doc.add_paragraph()

# Risiko 5
add_body("Risiko 5: Blind Spot — Tidak Ada Security Monitoring (CRITICAL)", bold=True, color=ACCENT_RED, size=Pt(12))
add_body("Hanya uptime monitoring berarti perusahaan BUTA terhadap serangan yang sedang terjadi.")
add_body("Bagaimana Dieksploitasi:", bold=True)
add_bullet("Slow exfiltration (DNS tunneling) tanpa terdeteksi")
add_bullet("Dwell time rata-rata global: 204 hari (Mandiant M-Trends 2023)")
add_bullet("Tanpa log yang proper, incident response menjadi reaktif bukan proaktif")
add_body("Dampak: Serangan berjalan berbulan-bulan tanpa terdeteksi (dalam kasus ini: 2 bulan); data terus bocor.", bold=True, color=ACCENT_RED)

doc.add_page_break()

# ============================================================
# SOAL 2
# ============================================================
add_styled_heading("SOAL 2: Perbaikan yang Dapat Dilakukan (20 Poin)", level=2, color=ACCENT_RED)

add_info_box("PERTANYAAN", "Sebutkan dan jelaskan apa saja yang dapat diperbaiki dari pipeline. Untuk setiap perbaikan, sebutkan tools spesifik dan bagaimana implementasinya.", "E8F4FD", ACCENT_BLUE)

add_styled_heading("JAWABAN SOAL 2", level=3, color=ACCENT_GREEN)

# Perbaikan 1
add_body("Perbaikan 1: Implementasi Access Control & Identity Management", bold=True, color=ACCENT_BLUE, size=Pt(12))

add_table(
    ["Aspek", "Implementasi"],
    [
        ["Individual Accounts", "Setiap developer harus memiliki akun personal dengan password unik"],
        ["MFA", "Wajibkan Multi-Factor Authentication untuk akses Jenkins, GitHub, dan production"],
        ["RBAC", "Bagi akses sesuai peran: Dev (read/build), Lead (merge/deploy), Admin (config)"],
        ["Auto Offboarding", "Saat resign, akses otomatis dicabut melalui Identity Provider (IdP)"],
        ["Least Privilege", "Developer tidak boleh memiliki akses root ke production"],
    ],
    header_color="0F4C81"
)

add_body("Tools: Keycloak/Okta (IAM), HashiCorp Vault (secrets), GitHub Branch Protection, Jenkins RBAC Plugin", bold=True, color=NAVY)

doc.add_paragraph()

# Perbaikan 2
add_body("Perbaikan 2: Integrasi Security Scanning dalam Pipeline", bold=True, color=ACCENT_BLUE, size=Pt(12))

add_body("a) SAST — Tool: SonarQube + Semgrep", bold=True)
add_body("Dijalankan setiap push/PR. Mendeteksi SQL Injection, XSS, Command Injection, dll.")
add_code_block("stage('SAST Scan') {\n  steps {\n    sh 'sonar-scanner -Dsonar.projectKey=nusacloud-erp'\n    sh 'semgrep scan --config=auto --severity=ERROR'\n    waitForQualityGate abortPipeline: true\n  }\n}")

add_body("b) SCA — Tool: Snyk + OWASP Dependency-Check", bold=True)
add_body("Dijalankan setiap build + scheduled weekly. Mendeteksi CVE di dependencies.")
add_code_block("stage('SCA Scan') {\n  steps {\n    sh 'snyk test --severity-threshold=high'\n    sh 'dependency-check --project nusacloud --scan ./src'\n  }\n}")

add_body("c) Container Scanning — Tool: Trivy", bold=True)
add_body("Dijalankan setelah Docker build, sebelum push ke registry.")
add_code_block("stage('Container Scan') {\n  steps {\n    sh 'trivy image --severity HIGH,CRITICAL --exit-code 1 nusacloud-erp:${BUILD_NUMBER}'\n  }\n}")

add_body("d) DAST — Tool: OWASP ZAP", bold=True)
add_body("Dijalankan setelah deploy ke staging, sebelum production.")
add_code_block("stage('DAST Scan') {\n  steps {\n    sh 'docker-compose -f docker-compose.staging.yml up -d'\n    sh 'zap-cli quick-scan http://staging.nusacloud.internal:8080'\n    sh 'zap-cli report -o zap-report.html -f html'\n  }\n}")

doc.add_paragraph()

# Perbaikan 3
add_body("Perbaikan 3: Secret Management", bold=True, color=ACCENT_BLUE, size=Pt(12))
add_body("Tool: HashiCorp Vault + GitLeaks", bold=True)
add_bullet("Pre-commit hook dengan GitLeaks untuk mendeteksi secrets sebelum masuk ke repo")
add_bullet("Inject secrets dari Vault saat runtime, bukan hardcode di source code")
add_bullet("Rotasi secrets secara berkala dan otomatis")

doc.add_paragraph()

# Perbaikan 4
add_body("Perbaikan 4: Code Review & Approval Workflow", bold=True, color=ACCENT_BLUE, size=Pt(12))
add_bullet("Mandatory Pull Request — semua perubahan harus melalui PR")
add_bullet("Minimum 2 Approvers — minimal 2 reviewer meng-approve")
add_bullet("Automated Security Review — bot otomatis menjalankan security scan pada PR")
add_bullet("Signed Commits — GPG signed commits untuk keaslian kode")

doc.add_paragraph()

# Perbaikan 5
add_body("Perbaikan 5: Security Monitoring & Incident Response", bold=True, color=ACCENT_BLUE, size=Pt(12))
add_body("Tools: Wazuh (SIEM), ELK Stack (log management), Falco (runtime), Prometheus+Grafana (metrics)")
add_bullet("Centralized logging — semua log dikumpulkan di satu tempat")
add_bullet("Alert rules: unusual DNS, failed logins, data exfiltration patterns, pipeline changes")
add_bullet("24/7 monitoring dengan on-call rotation")

doc.add_page_break()

# ============================================================
# SOAL 3
# ============================================================
add_styled_heading("SOAL 3: Redesign Pipeline (20 Poin)", level=2, color=ACCENT_RED)

add_info_box("PERTANYAAN", "Redesign pipeline CI/CD dengan prinsip DevSecOps. Gambarkan pipeline baru beserta setiap stage, tools, dan security gates.", "E8F4FD", ACCENT_BLUE)

add_styled_heading("JAWABAN SOAL 3", level=3, color=ACCENT_GREEN)

add_body("Pipeline CI/CD Baru — DevSecOps Integrated (9 Stages)", bold=True, color=NAVY, size=Pt(13))

pipeline_stages = [
    ["Stage 1: PLAN & DESIGN", "Threat Modeling (STRIDE/DREAD)\nSecurity Requirements\nRisk Assessment", "Microsoft Threat Modeling Tool\nOWASP Threat Dragon", "Threat model approved"],
    ["Stage 2: CODE & COMMIT", "IDE Security Plugins\nPre-commit Hooks\nSigned Commits\nBranch Protection", "GitLeaks, Semgrep\nSonarLint, GPG", "No secrets detected\nLint pass"],
    ["Stage 3: PULL REQUEST", "Mandatory 2 Reviewer Approval\nAutomated Security Review\nPR Security Checklist", "GitHub Advanced Security\nCodeQL", "2 approvals\nSecurity check pass"],
    ["Stage 4: BUILD & SAST/SCA", "Compile & Build\nSAST Scan\nSCA Scan\nUnit Tests", "SonarQube, Semgrep\nSnyk, OWASP Dep-Check", "0 Critical/High findings"],
    ["Stage 5: CONTAINER", "Docker Image Build\nContainer Scan\nDockerfile Lint\nImage Signing", "Trivy, Hadolint\nCosign", "No HIGH CVE\nImage signed"],
    ["Stage 6: STAGING & DAST", "Deploy ke Staging\nDAST Scan\nIntegration Tests\nIaC Scanning", "OWASP ZAP\nCheckov", "DAST pass\nNo critical findings"],
    ["Stage 7: APPROVAL", "Security Team Sign-off\nCompliance Check (UU PDP)\nChange Management", "Manual Process\nCompliance Framework", "Security team approved"],
    ["Stage 8: PRODUCTION", "Blue/Green atau Canary Deploy\nSecret Injection\nConfig Audit\nRollback Plan", "Terraform, HashiCorp Vault", "Deployment verified"],
    ["Stage 9: MONITOR", "SIEM\nLog Aggregation\nRuntime Protection\nAlerting & IRP", "Wazuh, ELK Stack\nFalco, PagerDuty", "Continuous 24/7"],
]

add_table(
    ["Stage", "Aktivitas", "Tools", "Security Gate"],
    pipeline_stages,
    header_color="0D1B2A"
)

add_info_box("ATURAN KRITIS", "Jika SATU SAJA security gate FAIL, deployment ke production DIHENTIKAN sampai masalah diperbaiki. Ini adalah prinsip 'Break the Build'.", "FDEDEC", ACCENT_RED)

add_styled_heading("Summary Security Gates:", level=3, color=NAVY)

add_code_block(
    "Gate 1 (Code):      Tidak ada secrets                    → PASS/FAIL\n"
    "Gate 2 (Review):     2 Reviewer approval                  → PASS/FAIL\n"
    "Gate 3 (Build):      SAST + SCA = 0 Critical/High         → PASS/FAIL\n"
    "Gate 4 (Container):  Trivy scan = 0 Critical              → PASS/FAIL\n"
    "Gate 5 (Staging):    DAST findings = 0 Critical           → PASS/FAIL\n"
    "Gate 6 (Approval):   Security team sign-off               → PASS/FAIL"
)

doc.add_page_break()

# ============================================================
# SOAL 4
# ============================================================
add_styled_heading("SOAL 4: What-If Scenario (20 Poin)", level=2, color=ACCENT_RED)

add_info_box("PERTANYAAN", "Bagaimana jika PT. NusaCloud sudah mengimplementasikan pipeline DevSecOps baru, tetapi seorang developer secara tidak sengaja memasukkan library open-source yang berisi kerentanan zero-day (belum ada CVE) — dan library tersebut diexploit 1 minggu setelah deployment?", "E8F4FD", ACCENT_BLUE)

add_styled_heading("JAWABAN SOAL 4", level=3, color=ACCENT_GREEN)

# Part a
add_body("a) Apakah Pipeline Baru Bisa Mencegah Insiden Ini Sepenuhnya?", bold=True, color=NAVY, size=Pt(12))

add_body("TIDAK, pipeline baru TIDAK BISA sepenuhnya mencegah insiden zero-day.", bold=True, color=ACCENT_RED)

add_body("Alasan:")
add_bullet("SAST hanya deteksi pola kode yang sudah diketahui berbahaya — zero-day belum ada pattern/rule")
add_bullet("SCA cocokkan versi library dengan database CVE — zero-day berarti belum ada CVE-nya")
add_bullet("Container Scanning juga bergantung pada database CVE yang sama")
add_bullet("DAST mungkin deteksi beberapa anomali, tapi tidak dijamin")

add_info_box("KESIMPULAN", "Tidak ada tool yang memberikan 100% protection. DevSecOps adalah defense-in-depth (pertahanan berlapis) dan reducing attack surface, bukan eliminasi total risiko.", "E8F8F5", ACCENT_GREEN)

doc.add_paragraph()

# Part b
add_body("b) Stage yang Paling Relevan", bold=True, color=NAVY, size=Pt(12))

add_body("Stage 9: Monitoring & Response adalah stage paling kritis untuk skenario ini.", bold=True)

add_table(
    ["Komponen", "Peran dalam Deteksi Zero-Day"],
    [
        ["SIEM (Wazuh/Splunk)", "Deteksi anomali behavior: unusual traffic, unexpected processes, privilege escalation"],
        ["Runtime Protection (Falco)", "Deteksi perilaku mencurigakan di container saat runtime"],
        ["Log Aggregation (ELK)", "Forensic analysis cepat + retention min. 90 hari"],
        ["Stage 6 - DAST", "Mungkin deteksi behavior anomali jika exploit menghasilkan response abnormal"],
    ],
    header_color="007B83"
)

doc.add_paragraph()

# Part c
add_body("c) Langkah Mitigasi Tambahan", bold=True, color=NAVY, size=Pt(12))

add_table(
    ["#", "Mitigasi", "Detail"],
    [
        ["1", "RASP", "Runtime Application Self-Protection (Contrast Security) — blokir serangan real-time"],
        ["2", "WAF", "Web Application Firewall (ModSecurity, Cloudflare WAF) — filter malicious requests"],
        ["3", "Network Segmentation", "Micro-segments untuk mencegah lateral movement"],
        ["4", "UEBA", "User & Entity Behavior Analytics — deteksi anomali berdasarkan baseline"],
        ["5", "Incident Response Plan", "Playbook siap: siapa bertanggung jawab, cara isolasi, komunikasi ke klien"],
        ["6", "Canary Deployment", "Deploy ke 5-10% user dulu, monitor 24-48 jam, baru full rollout"],
        ["7", "Bug Bounty Program", "External researchers menemukan kerentanan sebelum attacker"],
        ["8", "Threat Intelligence", "Subscribe ke feed (AlienVault OTX, MISP) untuk info zero-day tercepat"],
    ],
    header_color="0F4C81"
)

add_styled_heading("Timeline Response yang Ideal:", level=3, color=NAVY)

add_code_block(
    "Zero-day exploit terjadi           (T+0)\n"
    "SIEM/Falco mendeteksi anomali      (T+5 menit)\n"
    "Alert ke Security Team             (T+6 menit)\n"
    "Investigation dimulai              (T+15 menit)\n"
    "Keputusan: Isolasi atau Rollback   (T+30 menit)\n"
    "Rollback ke versi sebelumnya       (T+45 menit)\n"
    "Post-incident review               (T+24 jam)"
)

add_info_box("PERBANDINGAN", "Pipeline lama: insiden terdeteksi setelah 2 BULAN. Pipeline DevSecOps: respons dalam hitungan MENIT hingga JAM.", "E8F8F5", ACCENT_GREEN)

doc.add_page_break()

# ============================================================
# SOAL 5
# ============================================================
add_styled_heading("SOAL 5: Critical Missing Process (20 Poin)", level=2, color=ACCENT_RED)

add_info_box("PERTANYAAN", "Identifikasi SATU proses kritikal yang hilang dari pipeline original. Jelaskan: a) Apa prosesnya dan mengapa penting; b) Bagaimana ketidakhadirannya berkontribusi pada insiden; c) Bagaimana implementasinya; d) Contoh output.", "E8F4FD", ACCENT_BLUE)

add_styled_heading("JAWABAN SOAL 5", level=3, color=ACCENT_GREEN)

add_body("Proses Kritikal yang Hilang: THREAT MODELING", bold=True, color=DARK_BLUE, size=Pt(14))

# Part a
add_body("a) Apa Itu Threat Modeling dan Mengapa Penting?", bold=True, color=NAVY, size=Pt(12))

add_body("Threat Modeling adalah proses terstruktur untuk mengidentifikasi, mengevaluasi, dan memprioritaskan potensi ancaman keamanan SEBELUM sistem dibangun. Ini dilakukan di tahap paling awal (Plan/Design).")

add_body("Mengapa sangat penting:", bold=True)
add_bullet("Proaktif vs Reaktif — identifikasi masalah SEBELUM kode ditulis (6x lebih murah)")
add_bullet("Holistic View — melihat arsitektur, data flow, trust boundaries, proses operasional")
add_bullet("Risk Prioritization — prioritaskan mitigasi berdasarkan likelihood dan impact")
add_bullet("Compliance — UU PDP dan ISO 27001 mensyaratkan pendekatan risk-based")

add_body("Framework Threat Modeling:", bold=True)
add_table(
    ["Framework", "Deskripsi"],
    [
        ["STRIDE", "Spoofing, Tampering, Repudiation, Information Disclosure, DoS, Elevation of Privilege"],
        ["DREAD", "Damage, Reproducibility, Exploitability, Affected Users, Discoverability — risk scoring"],
        ["PASTA", "Process for Attack Simulation and Threat Analysis — risk-centric"],
        ["LINDDUN", "Fokus privacy threats — relevan untuk UU PDP"],
    ],
    header_color="0F4C81"
)

doc.add_paragraph()

# Part b
add_body("b) Bagaimana Ketidakhadiran Threat Modeling Berkontribusi pada Insiden?", bold=True, color=NAVY, size=Pt(12))

add_body("Analisis STRIDE terhadap Pipeline NusaCloud — SEMUA kategori dilanggar:", bold=True)

add_table(
    ["Threat Category", "Threat yang Seharusnya Teridentifikasi", "Status"],
    [
        ["Spoofing", "Mantan karyawan bisa menyamar sebagai admin (shared creds)", "❌ Tidak teridentifikasi"],
        ["Tampering", "Jenkinsfile bisa dimodifikasi tanpa kontrol", "❌ Tidak teridentifikasi"],
        ["Repudiation", "Tidak ada audit log — pelaku tidak bisa diidentifikasi", "❌ Tidak teridentifikasi"],
        ["Info Disclosure", "Hardcoded secrets, data flow tidak terenkripsi", "❌ Tidak teridentifikasi"],
        ["Denial of Service", "Single point of failure di Jenkins dan production", "❌ Tidak teridentifikasi"],
        ["Elevation of Privilege", "Semua developer memiliki root access", "❌ Tidak teridentifikasi"],
    ],
    header_color="C0392B"
)

doc.add_paragraph()

# Part c
add_body("c) Bagaimana Threat Modeling Diimplementasikan?", bold=True, color=NAVY, size=Pt(12))

add_body("Langkah 1: Identifikasi Assets", bold=True)
add_code_block(
    "Asset Register untuk NusaCloud ERP:\n"
    "├── Data Assets\n"
    "│   ├── Client Financial Data (CRITICAL)\n"
    "│   ├── Employee PII (CRITICAL)\n"
    "│   ├── Transaction Records (HIGH)\n"
    "│   └── Audit/Tax Reports (HIGH)\n"
    "├── System Assets\n"
    "│   ├── Production Server (CRITICAL)\n"
    "│   ├── Jenkins Build Server (HIGH)\n"
    "│   └── Docker Registry (MEDIUM)\n"
    "└── Credential Assets\n"
    "    ├── Database Credentials (CRITICAL)\n"
    "    ├── API Keys (HIGH)\n"
    "    └── SSL/TLS Certificates (HIGH)"
)

add_body("Langkah 2: Buat Data Flow Diagram (DFD)", bold=True)
add_bullet("Identifikasi semua komponen sistem dan aliran data")
add_bullet("Tentukan trust boundaries")
add_bullet("Identifikasi entry points dan exit points")

add_body("Langkah 3: Identifikasi Threats (STRIDE)", bold=True)
add_bullet("Untuk setiap data flow melewati trust boundary, analisis ke-6 kategori STRIDE")

add_body("Langkah 4: Risk Scoring (DREAD)", bold=True)
add_body("Contoh scoring — Threat: 'Mantan karyawan mengakses Jenkins':", italic=True)

add_table(
    ["DREAD Factor", "Score (1-10)", "Justification"],
    [
        ["Damage", "10", "Full access production, data breach masif"],
        ["Reproducibility", "9", "Mudah diulang karena shared credentials"],
        ["Exploitability", "10", "Hanya perlu login biasa, no exploit needed"],
        ["Affected Users", "10", "500+ perusahaan klien, jutaan end-users"],
        ["Discoverability", "8", "Shared credentials = common knowledge"],
        ["TOTAL", "47/50", "CRITICAL — Harus dimitigasi segera"],
    ],
    header_color="C0392B"
)

add_body("Langkah 5: Tentukan Mitigasi", bold=True)
add_table(
    ["Threat", "DREAD", "Mitigasi", "Prioritas"],
    [
        ["Unauthorized access (shared creds)", "47/50", "Individual accounts + MFA + RBAC", "P1 - Immediate"],
        ["Pipeline tampering", "42/50", "Branch protection + signed commits + audit", "P1 - Immediate"],
        ["Hardcoded secrets", "40/50", "HashiCorp Vault + GitLeaks", "P1 - Immediate"],
        ["Unscanned containers", "38/50", "Trivy scanning + image signing", "P2 - Short-term"],
        ["No security monitoring", "45/50", "SIEM (Wazuh) + ELK + Falco", "P1 - Immediate"],
    ],
    header_color="0F4C81"
)

add_body("Langkah 6: Review dan Update", bold=True)
add_bullet("Review setiap perubahan arsitektur")
add_bullet("Minimal quarterly review")
add_bullet("Update setelah setiap insiden keamanan")

doc.add_paragraph()

# Part d
add_body("d) Contoh Output Threat Modeling Document", bold=True, color=NAVY, size=Pt(12))

add_code_block(
    "════════════════════════════════════════════════\n"
    "   THREAT MODEL REPORT\n"
    "   PT. NusaCloud Technologies — ERP Platform\n"
    "   Version: 1.0 | Date: Q1 2026\n"
    "════════════════════════════════════════════════\n"
    "\n"
    "EXECUTIVE SUMMARY:\n"
    "Ditemukan 15 threats:\n"
    "  6 CRITICAL | 4 HIGH | 3 MEDIUM | 2 LOW\n"
    "Average DREAD Score: 35.2/50\n"
    "Highest: 47/50 (Unauthorized Pipeline Access)\n"
    "\n"
    "RECOMMENDATION:\n"
    "Hentikan deployment sampai mitigasi\n"
    "6 CRITICAL threats diimplementasikan.\n"
    "\n"
    "[TM-001] Unauthorized Build Pipeline Access\n"
    "  Category: Spoofing + Elevation of Privilege\n"
    "  DREAD: 47/50\n"
    "  Mitigation: Individual accounts, MFA, RBAC\n"
    "  Owner: DevOps Lead | Deadline: Sprint 1\n"
    "\n"
    "[TM-002] Supply Chain Injection\n"
    "  Category: Tampering\n"
    "  DREAD: 42/50\n"
    "  Mitigation: Snyk + OWASP Dep-Check\n"
    "  Owner: Security Team | Deadline: Sprint 1\n"
)

doc.add_page_break()

# ============================================================
# BAGIAN C: QUICK REFERENCE
# ============================================================
add_styled_heading("BAGIAN C: Quick Reference — Tools DevSecOps", level=1, color=DARK_BLUE)
add_separator()

add_table(
    ["Kategori", "Tools"],
    [
        ["Planning", "OWASP Threat Dragon, Microsoft TMT"],
        ["Secret Detection", "GitLeaks, TruffleHog"],
        ["SAST", "SonarQube, Semgrep, Checkmarx, CodeQL"],
        ["SCA", "Snyk, Trivy, OWASP Dependency-Check, Dependabot"],
        ["Container Scan", "Trivy, Aqua Security, Grype, Docker Scout"],
        ["IaC Scanning", "Checkov, Terrascan, tfsec, KICS"],
        ["DAST", "OWASP ZAP, Burp Suite, Nikto, Nuclei"],
        ["Secret Management", "HashiCorp Vault, AWS Secrets Manager"],
        ["IAM", "Keycloak, Okta"],
        ["SIEM/Monitoring", "Wazuh, Splunk, ELK Stack"],
        ["Runtime Protection", "Falco, Aqua Security"],
        ["CI/CD Platform", "Jenkins, GitLab CI, GitHub Actions"],
        ["WAF", "ModSecurity, AWS WAF, Cloudflare"],
    ],
    header_color="0D1B2A"
)

doc.add_page_break()

# ============================================================
# BAGIAN D: TIPS UJIAN
# ============================================================
add_styled_heading("BAGIAN D: Tips Menjawab Ujian", level=1, color=DARK_BLUE)
add_separator()

add_styled_heading("Strategi Menjawab Study Case DevSecOps", level=2, color=ACCENT_BLUE)

tips = [
    ("1. Selalu mulai dengan identifikasi masalah", "Sebutkan masalah spesifik dari case sebelum memberikan solusi"),
    ("2. Gunakan framework", "STRIDE, DREAD, Defense-in-Depth, Shift-Left"),
    ("3. Sebutkan tools spesifik", "Dosen ingin dengar nama tools, bukan hanya konsep"),
    ("4. Jelaskan MENGAPA, bukan hanya APA", "Jangan hanya list tools, jelaskan kenapa relevan"),
    ("5. Hubungkan dengan real-world incidents", "SolarWinds, Log4Shell, CodeCov, dll."),
    ("6. Gambar diagram jika redesign", "Visual representation sangat penting"),
    ("7. Ingat: tidak ada security 100%", "Jawaban yang baik mengakui limitasi + defense-in-depth"),
]

for title, desc in tips:
    add_body(title, bold=True, color=ACCENT_BLUE)
    add_body(f"  → {desc}", color=MEDIUM_GRAY)

doc.add_paragraph()

add_styled_heading("Key Terms yang Harus Diingat", level=2, color=ACCENT_BLUE)

add_table(
    ["Term", "Definisi"],
    [
        ["Shift Left", "Menggeser security testing ke tahap awal development"],
        ["Defense in Depth", "Pertahanan berlapis, tidak bergantung pada satu kontrol"],
        ["Zero Trust", "Jangan pernah trust, selalu verify — bahkan internal users"],
        ["Security Gate", "Checkpoint otomatis sebelum lanjut ke stage berikutnya"],
        ["Supply Chain Attack", "Serangan melalui komponen third-party"],
        ["Dwell Time", "Waktu attacker dalam sistem tanpa terdeteksi"],
        ["Blast Radius", "Seberapa besar dampak jika satu komponen terkompromi"],
        ["SBOM", "Software Bill of Materials — daftar semua komponen software"],
        ["CVE", "Common Vulnerabilities and Exposures — database kerentanan publik"],
        ["SIEM", "Security Information and Event Management"],
    ],
    header_color="0D1B2A"
)

doc.add_paragraph()
doc.add_paragraph()

add_separator()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dokumen ini dibuat untuk persiapan UAS Server Network Administration")
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.italic = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BINUS University — Cyber Security Program — Semester 4")
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.italic = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Last Updated: April 2026")
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.italic = True
run.font.name = 'Calibri'

# ============================================================
# SAVE
# ============================================================
output_path = r"c:\SEMESTER4\SERVER_NETWORK_ADMIN\UAS_DEVSECOPS_STUDY_CASE.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
